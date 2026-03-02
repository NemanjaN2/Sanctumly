"""
Document upload routes for RAG
"""

import io
import json
import csv
import logging
from fastapi import APIRouter, HTTPException, UploadFile, File, Depends
from sqlalchemy.orm import Session

import PyPDF2
import docx
import openpyxl
import xlrd
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup

from app.database import get_db
from app.models.document import UserDocument
from app.security import sanitize_session_id
from app.services.rag import generate_embedding, chunk_text
from app.config import MAX_FILE_SIZE

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/upload", tags=["Upload"])


@router.post("/document")
async def upload_document(
    file: UploadFile = File(...),
    session_id: str = "",
    db: Session = Depends(get_db)
):
    """Upload and process document with RAG - Supports PDF, TXT, DOCX, XLSX, XLS, CSV, EPUB"""
    try:
        session_id = sanitize_session_id(session_id)
        
        content = await file.read()
        
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size is {MAX_FILE_SIZE / (1024*1024):.0f}MB"
            )
        
        text_content = ""
        
        # Extract text based on file type
        if file.filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                text_content += page.extract_text()
        
        elif file.filename.endswith('.txt'):
            text_content = content.decode('utf-8')
        
        elif file.filename.endswith('.docx'):
            doc = docx.Document(io.BytesIO(content))
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
        
        elif file.filename.endswith('.xlsx'):
            wb = openpyxl.load_workbook(io.BytesIO(content))
            for sheet in wb.worksheets:
                text_content += f"Sheet: {sheet.title}\n"
                for row in sheet.iter_rows(values_only=True):
                    text_content += ' | '.join([str(cell) if cell else '' for cell in row]) + "\n"
                text_content += "\n"
        
        elif file.filename.endswith('.xls'):
            wb = xlrd.open_workbook(file_contents=content)
            for sheet in wb.sheets():
                text_content += f"Sheet: {sheet.name}\n"
                for row_idx in range(sheet.nrows):
                    text_content += ' | '.join([str(cell.value) for cell in sheet.row(row_idx)]) + "\n"
                text_content += "\n"
        
        elif file.filename.endswith('.csv'):
            decoded_content = content.decode('utf-8')
            csv_reader = csv.reader(io.StringIO(decoded_content))
            for row in csv_reader:
                text_content += ' | '.join(row) + "\n"
        
        elif file.filename.endswith('.epub'):
            book = epub.read_epub(io.BytesIO(content))
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text_content += soup.get_text() + "\n"
        
        else:
            text_content = content.decode('utf-8', errors='ignore')
        
        chunks = chunk_text(text_content)
        
        logger.info(f"📄 Processing {file.filename}: {len(chunks)} chunks, {len(text_content)} chars")
        
        # Store chunks with embeddings
        for idx, chunk in enumerate(chunks):
            embedding = generate_embedding(chunk)
            
            doc = UserDocument(
                session_id=session_id,
                filename=file.filename,
                chunk_index=idx,
                content=chunk,
                embedding=json.dumps(embedding) if embedding else None
            )
            db.add(doc)
        
        db.commit()
        
        logger.info(f"✅ Uploaded {file.filename}: {len(chunks)} chunks stored")
        
        return {
            "success": True,
            "filename": file.filename,
            "chunks": len(chunks),
            "size": len(text_content)
        }
        
    except Exception as e:
        logger.error(f"❌ Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
